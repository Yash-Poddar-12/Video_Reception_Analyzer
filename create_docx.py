import docx

doc = docx.Document()

# Adding Title
doc.add_heading('Programming for Data Science (BCSE207L) - Course Project', 0)

# 1. Problem Statement
doc.add_heading('1. Problem Statement', level=1)
doc.add_paragraph(
    "Identify a real-world data science problem and implement an end-to-end "
    "pipeline in R (data acquisition, preprocessing, modeling, visualization, "
    "and insights generation)."
)

# 2. Dataset Requirements
doc.add_heading('2. Dataset Requirements', level=1)
p2 = doc.add_paragraph()
p2.add_run("Source:").bold = True
p2.add_run(" Public reliable sources (e.g., Gov portals, APIs, World Bank, UCI). ")
p2.add_run("Kaggle and toy datasets are strictly not allowed.").bold = True
p2.add_run("\nCriteria:\n")
p2.add_run("- Minimum 10,000 rows\n")
p2.add_run("- Minimum 8 attributes\n")
p2.add_run("- At least 2 data types (numeric + categorical)")

# 3. API Integration
doc.add_heading('3. API Integration', level=1)
doc.add_paragraph(
    "Integrate at least one API for data ingestion or enrichment. You must demonstrate:"
)
doc.add_paragraph("- Authentication & Pagination handling", style='List Bullet')
doc.add_paragraph("- Converting JSON/XML to R DataFrames", style='List Bullet')
doc.add_paragraph("- Logging API calls in Git commits", style='List Bullet')

# 4. R Programming
doc.add_heading('4. R Programming Requirements', level=1)
doc.add_paragraph("Data Preparation:", style='List Bullet')
doc.add_paragraph("  API extraction (httr/jsonlite), web scraping (rvest), cleaning, and feature engineering.")
doc.add_paragraph("EDA:", style='List Bullet')
doc.add_paragraph("  Summary statistics, correlation, outlier assessment, and at least 5 meaningful ggplot2 visualizations.")
doc.add_paragraph("Modelling:", style='List Bullet')
doc.add_paragraph("  Implement at least one model (Regression, Classification, Clustering, Time-series, etc.) with relevant evaluation metrics (RMSE, accuracy, F1-score, etc.).")

# 5. Power BI
doc.add_heading('5. Power BI (with R Visuals)', level=1)
doc.add_paragraph(
    "Use R scripts inside Power BI to render at least two advanced custom visuals "
    "(ggplot2-based) within a dynamic interactive dashboard showing key stats, trends, and model predictions."
)

# 6. Version Control
doc.add_heading('6. GitHub Requirement', level=1)
p6 = doc.add_paragraph()
p6.add_run("- Minimum 10 meaningful commits across 2 branches (main & development)\n")
p6.add_run("- Detailed README explaining installation and project structure\n")
p6.add_run("- Ensure API keys are NOT pushed (use .gitignore)").bold = True

# 7. Docker
doc.add_heading('7. Docker Containerization', level=1)
doc.add_paragraph(
    "Publish the containerized application (excluding Power BI) "
    "to your Docker Hub account with proper name and version tags."
)

# 8. Evaluation & General Info
doc.add_heading('8. General Information', level=1)
doc.add_paragraph("- Team Size: Max 2 students per group")
doc.add_paragraph("- Evaluation Scheme (20 Marks Total):")
doc.add_paragraph("  • Problem Definition & Dataset Quality: 4 marks")
doc.add_paragraph("  • API Integration & Data Engineering: 4 marks")
doc.add_paragraph("  • R Programming & Containerization: 6 marks")
doc.add_paragraph("  • Power BI with R Visuals: 4 marks")
doc.add_paragraph("  • GitHub & README Documentation: 2 marks")

doc.save('Project_Requirements.docx')
print("File Project_Requirements.docx created successfully!")
